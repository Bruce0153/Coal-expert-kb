from __future__ import annotations

from pathlib import Path
from typing import List

from langchain_core.documents import Document

from coal_kb.loaders.base import BaseLoader, detect_language, normalize_text
from coal_kb.loaders.registry import register_loader


class DocxLoader(BaseLoader):
    def __init__(self) -> None:
        try:
            import docx  # noqa: F401
        except Exception as exc:  # pragma: no cover - optional dependency
            raise RuntimeError("python-docx is required for DocxLoader") from exc

    def load(self, path: str) -> List[Document]:
        import docx

        doc = docx.Document(path)
        docs: List[Document] = []
        for idx, para in enumerate(doc.paragraphs):
            content = normalize_text(para.text or "")
            if not content:
                continue
            lang = detect_language(content)
            docs.append(
                Document(
                    page_content=content,
                    metadata={
                        "source_file": path,
                        "section": "body",
                        "line_range": [idx + 1, idx + 1],
                        "doc_type": "docx",
                        "language": lang,
                        "parser": "docx",
                    },
                )
            )
        return docs


register_loader("docx", DocxLoader)
